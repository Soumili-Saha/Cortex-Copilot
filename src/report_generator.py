
"""
src/report_generator.py
Builds downloadable .docx reports, saved ONLY into the active tenant's
output_reports/ folder (tenant isolation preserved).

  * generate_anomaly_report(...)  -> Feature 2
  * generate_data_report(...)     -> Feature 3
"""
from __future__ import annotations
import logging
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src import analytics, chart_utils, config, llm, prompts

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("report_generator")


def _ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


#anomaly report
def generate_anomaly_report(tenant_display: str, start, end) -> Path:
    t = config.get_tenant(tenant_display)
    reports_dir: Path = t["reports"]
    reports_dir.mkdir(parents=True, exist_ok=True)

    df = analytics.load_tenant_df(tenant_display)
    window = analytics.filter_by_range(df, start, end)
    stats = analytics.univariate_stats(window)
    table_rows = analytics.build_anomaly_table(stats)
    top_cols = analytics.top_univariate_columns(stats, n=5)
    labels, scores, feat_cols = analytics.isolation_forest_anomalies(window)

    tcol, x = analytics.get_full_timestamp(window)
    if x is None:
        x = window.index

    doc = Document()

    
    title = doc.add_heading("Anomaly Detection Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        f"Tenant: {t['display']}",
        f"Time range: {start}  to  {end}",
        f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}",
        f"Readings analysed: {len(window)}",
    ):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    doc.add_section(WD_SECTION.NEW_PAGE)
    title_sectPr = doc.sections[0]._sectPr
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    title_sectPr.append(vAlign)

    tmp_imgs: list[Path] = []

    
    import numpy as _np
    labels_arr = _np.asarray(labels)
    anom_pos = _np.where(labels_arr == -1)[0]

    doc.add_heading("Normal vs Anomaly Trends (Top 5 Variables)", level=1)
    doc.add_paragraph(
        "For each variable: the top panel is the value over the selected time "
        "window; the bottom panel marks each reading as Normal or Anomalous "
        "per the multivariate Isolation Forest, so anomalous stretches are "
        "easy to spot against the trend above them.")
    for col in top_cols:
        if col not in window.columns:
            continue
        series = pd.to_numeric(window[col], errors="coerce")
        fig, (ax1, ax2) = plt.subplots(
            2, 1, figsize=(7, 3.6), sharex=True,
            gridspec_kw={"height_ratios": [2, 1]})

        ax1.plot(x, series, linewidth=1, color="#0f6b52", label="Normal")
        ax1.set_title(f"{col}: normal vs anomaly")
        ax1.set_ylabel(col)
        ax1.legend(loc="upper right", fontsize=8, frameon=False)
        ax1.grid(alpha=0.25)

        
        col_anom_idx = set(stats.get(col, {}).get("anomaly_index") or [])
        series_reset = series.reset_index(drop=True)
        if col_anom_idx:
            normal_vals = series_reset[~series_reset.index.isin(col_anom_idx)].dropna()
        else:
            normal_vals = series_reset.dropna()
        if len(normal_vals) < 5:  # not enough clean points, fall back
            normal_vals = series_reset.dropna()

        y_min, y_max = normal_vals.quantile(0.01), normal_vals.quantile(0.99)
        margin = (y_max - y_min) * 0.1
        if margin == 0: # Fallback if the variable is literally a flat constant
            margin = abs(series.mean()) * 0.1 if series.mean() != 0 else 1
        ax1.set_ylim(y_min - margin, y_max + margin)
        

        indicator = pd.Series(1.0, index=series.index)
        if len(anom_pos):
            valid = [i for i in anom_pos if i < len(series)]
            indicator.iloc[valid] = -1.0
        ax2.plot(x, indicator, linewidth=1, color="#0f6b52", label="Normal", zorder=1)
        anom_idx = indicator.index[indicator == -1.0]
        if len(anom_idx):
            ax2.scatter([x[i] if hasattr(x, "__getitem__") else i for i in anom_idx],
                       indicator.loc[anom_idx], color="red", s=28,
                       label="Anomalous", zorder=5)
        ax2.set_yticks([-1, 1]); ax2.set_yticklabels(["Anomalous", "Normal"])
        ax2.set_ylim(-1.6, 1.6)
        ax2.legend(loc="upper right", fontsize=8, frameon=False)
        chart_utils.style_time_axis(ax2)

        #Dynamic Date/Time X-Axis Formatting ---
        import matplotlib.dates as mdates
        time_span = x.max() - x.min()
        if time_span <= pd.Timedelta(days=1):
            # Less than or equal to a day: show precise hours (e.g., 14:00)
            ax2.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M'))
        else:
            # More than a day: show Date and Month (e.g., 21-May)
            ax2.xaxis.set_major_formatter(mdates.DateFormatter('%d-%b'))
        

        fig.tight_layout()
        img = reports_dir / f".tmp_nva_{col}_{_ts()}.png"
        fig.savefig(img, dpi=110); plt.close(fig)
        tmp_imgs.append(img)
        doc.add_picture(str(img), width=Inches(6))

    doc.add_heading("Multivariate Isolation Forest", level=1)
    fig, ax = plt.subplots(figsize=(7, 3))
    colors = ["red" if l == -1 else "steelblue" for l in labels]
    ax.scatter(range(len(scores)), scores, c=colors, s=15)
    ax.set_title("Isolation Forest anomaly scores (red = anomaly)")
    ax.set_xlabel("reading index"); ax.set_ylabel("anomaly score")
    fig.tight_layout()
    iso_img = reports_dir / f".tmp_iso_{_ts()}.png"
    fig.savefig(iso_img, dpi=110); plt.close(fig)
    tmp_imgs.append(iso_img)
    doc.add_picture(str(iso_img), width=Inches(6))
    n_multi = int((labels == -1).sum())
    doc.add_paragraph(f"Multivariate anomalies flagged: {n_multi} of {len(labels)} "
                      f"readings across features {feat_cols}.")

  
    doc.add_heading("Anomaly Percentage per Attribute", level=1)

    # Build % of readings flagged as outliers (robust median/MAD z-score > 3)
    # for every column that has stats, sort descending, keep the top 15 
    
    n_readings = len(window)
    attr_pct = {}
    for col, s in stats.items():
        idx = s.get("anomaly_index")
        if idx is not None and n_readings > 0:
            attr_pct[col] = 100.0 * len(idx) / n_readings
    top_attrs = sorted(attr_pct.items(), key=lambda kv: kv[1], reverse=True)[:15]

    if top_attrs:
        names = [c for c, _ in top_attrs][::-1]  
        pcts = [p for _, p in top_attrs][::-1]

        fig, ax = plt.subplots(figsize=(7, 0.35 * len(top_attrs) + 1))
        ax.barh(names, pcts, color="#c0605a")
        ax.set_title("Univariate anomaly rate by attribute (top 15)")
        ax.set_xlabel("% of readings beyond robust z-score 3 (median/MAD)")
        fig.tight_layout()
        pct_img = reports_dir / f".tmp_pct_{_ts()}.png"
        fig.savefig(pct_img, dpi=110); plt.close(fig)
        tmp_imgs.append(pct_img)
        doc.add_picture(str(pct_img), width=Inches(6))
    else:
        doc.add_paragraph("No per-attribute anomaly-rate data available.")

    #Anomaly Data Extraction (Top 5 & Most Anomalous Points)
    import numpy as _np
    
    # 1. Multivariate: Find the absolute most anomalous timestamps (lowest scores from Isolation Forest)
    multivariate_anomalies = []
    if len(scores) > 0:
        most_anomalous_idx = _np.argsort(scores)[:3] # Top 3 most extreme multi-variate events
        for idx in most_anomalous_idx:
            if labels[idx] == -1: # Ensure it's actually classified as an anomaly
                try:
                    ts_val = x.iloc[idx] if hasattr(x, "iloc") else x[idx]
                    multivariate_anomalies.append(f"Timestamp: {ts_val}, Anomaly Score: {scores[idx]:.3f} (Lower = More Anomalous)")
                except Exception:
                    pass

    # 2. Univariate: Temporal clusters
    anomaly_clusters = {}
    for col, s in stats.items():
        if s.get("anomaly_present") and "anomaly_index" in s:
            for idx in s["anomaly_index"]:
                anomaly_clusters.setdefault(idx, []).append(col)
                
    critical_events = {idx: cols for idx, cols in anomaly_clusters.items() if len(cols) >= 4}
    correlation_note = ""
    if critical_events:
        correlation_note = "\n\n*** CRITICAL TEMPORAL CORRELATIONS (Simultaneous Robust Z-Score > 3 Deviations) ***\n"
        for idx, cols in list(critical_events.items())[:3]:
            try:
                ts_val = x.iloc[idx] if hasattr(x, "iloc") else x[idx]
                col_vals = [f"{c} ({window[c].iloc[idx]:.2f})" for c in cols[:5]]
                correlation_note += f"- At exact timestamp {ts_val}, {len(cols)} variables deviated simultaneously: {', '.join(col_vals)}.\n"
            except Exception:
                pass

    # 3. Compile Facts for LLM
    facts_list = [
        "SYSTEM CONFIRMATION: The system used a robust, median/MAD-based z-score "
        "(threshold |z| > 3) for univariate anomaly detection, so a column's own "
        "outliers cannot inflate its threshold and mask other anomalies.",
        f"\nMULTIVARIATE MOST ANOMALOUS TIMES (Isolation Forest, {n_multi} total anomalies):"
    ]
    if multivariate_anomalies:
        facts_list.extend([f"- {m}" for m in multivariate_anomalies])
    else:
        facts_list.append("- No critical multivariate events detected.")
        
    facts_list.append("\nTOP 5 UNIVARIATE ANOMALIES (Most Extreme Point per Variable):")
    for col in top_cols: 
        s = stats[col]
        if s["anomaly_present"] and s.get("max_anomaly_idx") is not None:
            try:
                idx = s["max_anomaly_idx"]
                ts_val = x.iloc[idx] if hasattr(x, "iloc") else x[idx]
                val = s["max_anomaly_val"]
                z_score = s["max_anomaly_z"]
                facts_list.append(
                    f"- Variable: {col} | Most Anomalous Timestamp: {ts_val} | Value: {val:.2f} | Z-Score: {z_score:.2f}"
                )
            except Exception:
                pass

    facts = "\n".join(facts_list) + correlation_note

    insight = llm.generate_response(
        prompts.ANOMALY_INSIGHT_TEMPLATE.format(
            tenant=t["display"], start=start, end=end, facts=facts,
            glossary=prompts.ELECTRICAL_TERM_GLOSSARY),
        system=prompts.ANOMALY_INSIGHT_SYSTEM, temperature=0.2, max_tokens=1500,
    )
    
    for para in insight.split("\n"):
        para = para.strip()
        if not para:
            continue
        
        para = para.replace("**", "").replace("#", "") # Strip markdown
        
        if "CRITICAL INCIDENT DIAGNOSIS" in para or "SECONDARY ANOMALIES" in para or "TARGETED RECOMMENDATIONS" in para:

            doc.add_heading(para.replace(":", ""), level=2) 
        else:
            doc.add_paragraph(para)

    out_path = reports_dir / f"anomaly_report_{t['id']}_{_ts()}.docx"
    doc.save(out_path)

    for img in tmp_imgs:     
        try: img.unlink()
        except Exception: pass

    logger.info("Anomaly report saved -> %s", out_path)
    return out_path

def generate_data_report(tenant_display: str, subset: pd.DataFrame,
                         start, end, variables: list[str]) -> Path:
    t = config.get_tenant(tenant_display)
    reports_dir: Path = t["reports"]
    reports_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Data Retrieval Report", level=0)
    doc.add_paragraph(f"Tenant: {t['display']}")
    doc.add_paragraph(f"Time range: {start}  to  {end}")
    doc.add_paragraph(f"Variables: {', '.join(variables) if variables else 'all'}")
    doc.add_paragraph(f"Rows: {len(subset)}")
    doc.add_paragraph(f"Generated: {datetime.now():%Y-%m-%d %H:%M:%S}")

    cols = list(subset.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light List Accent 1"
    for j, c in enumerate(cols):
        table.rows[0].cells[j].text = str(c)
    # cap at a sane number of rows for the docx while noting the full count
    max_rows = 1000
    for _, row in subset.head(max_rows).iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(cols):
            cells[j].text = str(row[c])
    if len(subset) > max_rows:
        doc.add_paragraph(f"(Showing first {max_rows} of {len(subset)} rows.)")

    out_path = reports_dir / f"data_report_{t['id']}_{_ts()}.docx"
    doc.save(out_path)
    logger.info("Data report saved -> %s", out_path)
    return out_path
